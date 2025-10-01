import torch
import os
from dotenv import load_dotenv
from transformers import AutoModelForCausalLM, AutoTokenizer
from peft import PeftModel
import sys
import time
import argparse
from pathlib import Path
from typing import Optional, List

try:
    # Optional dependency for reading .docx
    from docx import Document  # python-docx
except Exception:
    Document = None

# --- CONFIGURATION ---
# Load environment variables (model paths)
load_dotenv()

BASE_MODEL_PATH = os.getenv("EUROLLM_MODEL_PATH")
FINETUNED_MODEL_PATH = os.getenv("EUROLLM_LORA_ADAPTER")
MARKER = "### Answer:"

# Check if model paths are set
if not BASE_MODEL_PATH or not FINETUNED_MODEL_PATH:
    print("❌ Error: Make sure you have set EUROLLM_MODEL_PATH and EUROLLM_LORA_ADAPTER in your .env file")
    sys.exit(1)

# ==============================================================================
# ✅ EDIT HERE: Enter the sentences and language codes you want to translate
# ==============================================================================
SENTENCES_TO_TRANSLATE = [
    {
        "language": "RO",
        "text": "Travel and Culture: Traveling opens the door to new cultures, perspectives, and experiences that can change how we see the world. Visiting a foreign country allows us to taste unique foods, hear different languages, and witness traditions that have been passed down through generations. Whether you're walking through ancient ruins in Greece or enjoying a local festival in Portugal, travel offers lessons that no classroom ever could.",
    },
    {
        "language": "RO",
        "text": "Technology and Society: In recent years, technology has drastically reshaped the way we communicate, work, and learn. From smartphones to artificial intelligence, our daily routines are increasingly influenced by digital tools. While these advancements offer convenience and efficiency, they also raise important questions about privacy, security, and our dependence on machines.",
    },
    {
        "language": "RO",
        "text": "Nature and the Environment: The natural world is both beautiful and fragile. Forests, oceans, and wildlife are essential to life on Earth, yet human activity has led to pollution, climate change, and habitat destruction. Protecting the environment is not just the responsibility of governments or scientists—everyone can contribute by making sustainable choices and raising awareness about the importance of conservation.",
    },
]
# ==============================================================================


@torch.no_grad()
def load_model_and_tokenizer():
    """Loads the model and tokenizer once."""
    print("⏳ Loading model and tokenizer...")

    tok = AutoTokenizer.from_pretrained(BASE_MODEL_PATH)
    if tok.pad_token is None:
        tok.pad_token = tok.eos_token

    model = AutoModelForCausalLM.from_pretrained(
        BASE_MODEL_PATH,
        torch_dtype=torch.bfloat16,
        device_map="auto",
    )

    model = PeftModel.from_pretrained(model, FINETUNED_MODEL_PATH)
    model.eval()

    print("✅ Model loaded and ready!\n")
    return tok, model

def generate_translation(model, tokenizer, src_text, target_lang):
    """Generates a single translation."""
    instruction = f"Translate the following English text to {target_lang}: '{src_text}'"
    prompt = f"{instruction}\n{MARKER}\n"

    inputs = tokenizer(prompt, return_tensors="pt").to(model.device)

    outputs = model.generate(
        **inputs,
        max_new_tokens=150,
        temperature=0.1,
        do_sample=True,
        eos_token_id=tokenizer.eos_token_id,
        pad_token_id=tokenizer.pad_token_id,
    )

    full_output = tokenizer.decode(outputs[0], skip_special_tokens=True)
    prediction = full_output[len(prompt):].strip()

    return prediction

def read_docx_paragraphs(docx_path: Path) -> List[str]:
    """Extract non-empty paragraphs from a .docx file.

    Requires python-docx. If missing, instructs user to install it.
    """
    if Document is None:
        raise RuntimeError(
            "python-docx non installato. Esegui: pip install python-docx"
        )
    if not docx_path.exists():
        raise FileNotFoundError(f"File .docx non trovato: {docx_path}")

    doc = Document(str(docx_path))
    paras = [p.text.strip() for p in doc.paragraphs]
    return [p for p in paras if p]

def build_jobs_from_docx(docx_path: Path, target_lang: str) -> List[dict]:
    paragraphs = read_docx_paragraphs(docx_path)
    return [{"language": target_lang, "text": p} for p in paragraphs]


def parse_args():
    parser = argparse.ArgumentParser(description="Translate text or a .docx document.")
    parser.add_argument(
        "--docx",
        type=str,
        default=None,
        help="Percorso al file .docx da tradurre (es. resources/input.docx)",
    )
    parser.add_argument(
        "--lang",
        type=str,
        default=None,
        help="Codice lingua di destinazione (es. IT, RO, FR).",
    )
    parser.add_argument(
        "--out-docx",
        type=str,
        default=None,
        help="Percorso del file .docx di output con le traduzioni.",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    tokenizer, model = load_model_and_tokenizer()

    # Determine translation jobs
    jobs: List[dict]
    target_lang = args.lang
    docx_input_path: Optional[Path] = None

    if args.docx:
        docx_path = Path(args.docx)
        if target_lang is None:
            target_lang = "IT"
        try:
            jobs = build_jobs_from_docx(docx_path, target_lang)
            print(f"Caricato file .docx: {docx_path} ({len(jobs)} paragrafi)")
            docx_input_path = docx_path
        except FileNotFoundError:
            print("File .docx non trovato. Uso gli esempi hardcoded.")
            jobs = SENTENCES_TO_TRANSLATE
            docx_input_path = None
    else:
        # Fallback to hardcoded examples
        jobs = SENTENCES_TO_TRANSLATE
        print("Nessun --docx passato. Uso gli esempi hardcoded.")

    if target_lang is None and jobs and "language" in jobs[0]:
        target_lang = jobs[0]["language"]

    print("--- Starting Translations ---")
    total_start = time.perf_counter()

    translations: List[str] = []
    for i, job in enumerate(jobs):
        tgt = job["language"]
        src = job["text"]

        print(f"\n--- Translation #{i+1} ---")
        print(f"Original (-> {tgt}): {src}")

        print("Translating...")
        t0 = time.perf_counter()
        translated_text = generate_translation(model, tokenizer, src, tgt)
        t1 = time.perf_counter()

        print(f"Translation: {translated_text}")
        print(f"Duration: {format_mm_ss(t1 - t0)}")
        translations.append(translated_text)

    print("\n\nAll translations completed.")
    total_elapsed = time.perf_counter() - total_start
    print(f"Total time: {format_mm_ss(total_elapsed)}")


def format_mm_ss(seconds: float) -> str:
    """Format seconds into minutes and seconds, e.g., '2m 05s'."""
    total = int(round(seconds))
    m, s = divmod(total, 60)
    return f"{m}m {s:02d}s"

    if args.out_docx:
        out_docx_path = Path(args.out_docx)
    else:
        # Save in resources by default
        out_docx_path = Path("resources") / f"translations__{target_lang}.docx"

    # Write translations to .docx
    if out_docx_path is not None:
        write_translations_docx(translations, out_docx_path)
        print(f"Traduzioni salvate in: {out_docx_path}")


def write_translations_docx(paragraphs: List[str], out_path: Path) -> None:
    """Write translated paragraphs to a .docx file (one per paragraph)."""
    if Document is None:
        raise RuntimeError("python-docx non installato. Esegui: pip install python-docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(out_path))


if __name__ == "__main__":
    main()
